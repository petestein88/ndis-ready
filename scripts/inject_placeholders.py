#!/usr/bin/env python3
"""
NDIS Ready — Inject {{placeholders}} into all 65 template docs.

The 65 paid templates in Supabase Storage (templates/65 Files/) are
identical empty ~35KB skeletons. Per the product owner's instruction:
"for each doc, just paste in the fields into the empty doc that would be
needed in that particular doc ... just paste them in there."

This script BUILDS a clean .docx per document containing:
  - A title (the document name)
  - A document-control table (org identity + review dates + version)
  - The narrative {{placeholders}} relevant to that document category
    using docxtemplater-compatible {{ }} delimiters.

generate-documents.js merges these {{placeholders}} at delivery time via
docxtemplater, injecting the customer's real profile values.

Then it uploads each generated doc back to templates/65 Files/<name>.docx
using the Supabase Storage REST API (service_role key required).

Usage:
    SUPABASE_URL=... SUPABASE_SERVICE_KEY=... python3 inject_placeholders.py
"""

import os
import sys
import io
import time
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://ewfbovticpnbxxusekze.supabase.co").rstrip("/")
SERVICE_KEY  = os.environ.get("SUPABASE_SERVICE_KEY", "").strip()
BUCKET = "templates"
PREFIX = "65 Files"

# ---------------------------------------------------------------------------
# The 65 documents — id, name, category. Must match DOCUMENT_LIBRARY in
# api/generate-documents.js exactly.
# ---------------------------------------------------------------------------
DOCS = [
    ("1.1", "1.1 - Participant Rights Policy", "rights"),
    ("1.2", "1.2 - Complaints and Feedback Policy", "incidents"),
    ("1.3", "1.3 - Incident Management Policy", "incidents"),
    ("1.4", "1.4 - Reportable Incidents Procedure", "incidents"),
    ("1.5", "1.5 - Abuse and Neglect Prevention Policy", "safeguarding"),
    ("2.1", "2.1 - Governance Framework", "governance"),
    ("2.2", "2.2 - Risk Management Policy", "governance"),
    ("2.3", "2.3 - Financial Management Policy", "governance"),
    ("2.4", "2.4 - Human Resources Policy", "hr"),
    ("2.5", "2.5 - Staff Recruitment and Screening Procedure", "hr"),
    ("2.6", "2.6 - Staff Training and Development Policy", "hr"),
    ("2.7", "2.7 - Performance Management Procedure", "hr"),
    ("2.8", "2.8 - Whistleblower Policy", "governance"),
    ("2.9", "2.9 - Conflict of Interest Policy", "governance"),
    ("2.10", "2.10 - Record Keeping and Privacy Policy", "governance"),
    ("2.11", "2.11 - Information Management Policy", "governance"),
    ("2.12", "2.12 - Business Continuity Plan", "governance"),
    ("3.1", "3.1 - Service Delivery Policy", "support"),
    ("3.2", "3.2 - Intake and Eligibility Procedure", "support"),
    ("3.3", "3.3 - Support Planning Policy", "support"),
    ("3.4", "3.4 - Individual Support Plan Template", "support"),
    ("3.5", "3.5 - Person-Centred Practice Framework", "support"),
    ("3.6", "3.6 - Consent Policy", "rights"),
    ("3.7", "3.7 - Supported Decision Making Policy", "rights"),
    ("3.8", "3.8 - Transition and Exit Planning Procedure", "support"),
    ("3.9", "3.9 - Cultural Diversity and Inclusion Policy", "rights"),
    ("3.10", "3.10 - Restrictive Practices Policy", "safety"),
    ("3.11", "3.11 - Behaviour Support Policy", "safety"),
    ("3.12", "3.12 - Medication Management Policy", "safety"),
    ("3.13", "3.13 - Health and Medical Support Procedure", "safety"),
    ("3.14", "3.14 - Manual Handling Policy", "safety"),
    ("3.15", "3.15 - Emergency and Disaster Management Plan", "safety"),
    ("3.16", "3.16 - Mealtime Management Procedure", "safety"),
    ("4.1", "4.1 - SIL Service Agreement Template", "sil"),
    ("4.2", "4.2 - SIL Roster of Care Template", "sil"),
    ("4.3", "4.3 - SIL House Rules Template", "sil"),
    ("4.4", "4.4 - SIL Tenancy Support Policy", "sil"),
    ("4.5", "4.5 - SIL Daily Living Support Procedure", "sil"),
    ("4.6", "4.6 - Overnight and Sleepover Policy", "sil"),
    ("4.7", "4.7 - Household Budget Management Procedure", "sil"),
    ("4.8", "4.8 - Transition to SIL Procedure", "sil"),
    ("4.9", "4.9 - SDA and SIL Coordination Policy", "sil"),
    ("5.1", "5.1 - Code of Conduct", "hr"),
    ("5.2", "5.2 - Staff Handbook", "hr"),
    ("5.3", "5.3 - Position Description - Support Worker", "hr"),
    ("5.4", "5.4 - Position Description - Team Leader", "hr"),
    ("5.5", "5.5 - Position Description - Service Manager", "hr"),
    ("5.6", "5.6 - Onboarding Checklist", "hr"),
    ("5.7", "5.7 - Staff NDIS Worker Screening Checklist", "hr"),
    ("5.8", "5.8 - Volunteer Policy", "hr"),
    ("5.9", "5.9 - Contractor Management Policy", "hr"),
    ("6.1", "6.1 - Quality Management Framework", "quality"),
    ("6.2", "6.2 - Internal Audit Schedule", "quality"),
    ("6.3", "6.3 - Internal Audit Template", "quality"),
    ("6.4", "6.4 - Continuous Improvement Register", "quality"),
    ("6.5", "6.5 - Corrective Action Procedure", "quality"),
    ("6.6", "6.6 - NDIS Practice Standards Self-Assessment", "quality"),
    ("6.7", "6.7 - Audit Preparation Checklist", "quality"),
    ("7.1", "7.1 - Incident Report Form", "incidents"),
    ("7.2", "7.2 - Complaints Register", "incidents"),
    ("7.3", "7.3 - Risk Register", "governance"),
    ("7.4", "7.4 - Asset Register", "governance"),
    ("7.5", "7.5 - Training Register", "hr"),
    ("7.6", "7.6 - Participant Feedback Form", "incidents"),
    ("7.7", "7.7 - Worker Incident Statement Form", "incidents"),
]

# ---------------------------------------------------------------------------
# Placeholder field sets. Every variable key here MUST exist in
# generateVariables() in api/generate-documents.js so docxtemplater fills it.
# ---------------------------------------------------------------------------

# Identity block — appears in EVERY document's control table.
IDENTITY_FIELDS = [
    ("Organisation (legal name)", "org_name"),
    ("Trading name", "trading_name"),
    ("Organisation type", "org_type"),
    ("ABN", "abn"),
    ("ACN", "acn"),
    ("Address", "full_address"),
    ("Phone", "phone"),
    ("Email", "email"),
    ("Website", "website"),
]

# Document-control block — appears in EVERY document footer table.
CONTROL_FIELDS = [
    ("Document title", "document_title"),
    ("Document owner", "document_owner"),
    ("Version", "version"),
    ("Approved / review date", "review_date"),
    ("Next review date", "next_review_date"),
    ("Practice standard reference", "practice_standard_ref"),
]

# Narrative {{placeholders}} that every policy/procedure shares.
NARRATIVE_FIELDS = [
    ("Purpose", "purpose_statement"),
    ("Scope", "scope_statement"),
    ("Policy statement", "policy_statement"),
    ("Procedures", "procedures"),
    ("Roles and responsibilities", "roles_and_responsibilities"),
    ("Related documents", "related_documents"),
]

# Category-specific extra fields (people / registration / ops) injected on top.
CATEGORY_EXTRA = {
    "rights":       [("Safeguarding officer", "safeguarding_officer_name")],
    "incidents":    [("Safeguarding officer", "safeguarding_officer_name"),
                     ("After-hours contact", "after_hours_contact"),
                     ("Emergency contact name", "emergency_contact_name"),
                     ("Emergency contact phone", "emergency_contact_phone")],
    "safeguarding": [("Safeguarding officer", "safeguarding_officer_name"),
                     ("Works with children", "support_types")],
    "governance":   [("Director", "director_name"),
                     ("Director title", "director_title"),
                     ("Compliance officer", "compliance_officer_name"),
                     ("Insurance provider", "insurance_provider"),
                     ("Insurance policy number", "insurance_policy_number")],
    "hr":           [("Service manager", "service_manager_name"),
                     ("Compliance officer", "compliance_officer_name"),
                     ("WHS officer", "whs_officer_name"),
                     ("Approx. staff count", "staff_count")],
    "support":      [("Service manager", "service_manager_name"),
                     ("Support types", "support_types"),
                     ("Service areas", "service_areas"),
                     ("Operating hours", "operating_hours")],
    "safety":       [("Safeguarding officer", "safeguarding_officer_name"),
                     ("WHS officer", "whs_officer_name"),
                     ("Emergency contact name", "emergency_contact_name"),
                     ("Emergency contact phone", "emergency_contact_phone")],
    "sil":          [("Service manager", "service_manager_name"),
                     ("Support types", "support_types"),
                     ("Service areas", "service_areas"),
                     ("After-hours contact", "after_hours_contact"),
                     ("Operating hours", "operating_hours")],
    "quality":      [("Compliance officer", "compliance_officer_name"),
                     ("Audit pathway", "audit_pathway"),
                     ("Registration status", "registration_status"),
                     ("NDIS provider number", "ndis_provider_number"),
                     ("Registration groups", "registration_groups")],
}

# Forms & registers (7.x) are tabular — they take identity + control + the
# operational contacts rather than the long narrative blocks.
FORM_IDS = {"7.1", "7.2", "7.3", "7.4", "7.5", "7.6", "7.7"}

ACCENT = RGBColor(0x6B, 0x5B, 0x95)   # lavender/plum to match brand
INK    = RGBColor(0x2B, 0x2B, 0x2B)


def _set_cell(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_field_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for label, key in rows:
        r = table.add_row().cells
        _set_cell(r[0], label, bold=True, color=ACCENT)
        _set_cell(r[1], "{{" + key + "}}", color=INK)
    doc.add_paragraph("")
    return table


def build_doc(doc_id, name, category):
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    trun = title.add_run(name.split(" - ", 1)[-1] if " - " in name else name)
    trun.bold = True
    trun.font.size = Pt(20)
    trun.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    srun = sub.add_run("Document ID " + doc_id + " · NDIS Practice Standards 2021")
    srun.font.size = Pt(9)
    srun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Organisation identity
    h = doc.add_paragraph(); hr = h.add_run("Organisation Details")
    hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = ACCENT
    _add_field_table(doc, IDENTITY_FIELDS)

    # Key people / category-specific fields
    extra = CATEGORY_EXTRA.get(category, [])
    if extra:
        h = doc.add_paragraph(); hr = h.add_run("Key People & Details")
        hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = ACCENT
        _add_field_table(doc, extra)

    # Narrative blocks (policies/procedures only — not raw forms/registers)
    if doc_id not in FORM_IDS:
        for label, key in NARRATIVE_FIELDS:
            h = doc.add_paragraph(); hr = h.add_run(label)
            hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = ACCENT
            body = doc.add_paragraph()
            body.add_run("{{" + key + "}}").font.size = Pt(11)
            doc.add_paragraph("")

    # Document control
    h = doc.add_paragraph(); hr = h.add_run("Document Control")
    hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = ACCENT
    _add_field_table(doc, CONTROL_FIELDS)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def upload(name, data):
    path = f"{PREFIX}/{name}.docx"
    url = f"{SUPABASE_URL}/storage/v1/object/{BUCKET}/{requests.utils.quote(path)}"
    headers = {
        "Authorization": f"Bearer {SERVICE_KEY}",
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "x-upsert": "true",
    }
    resp = requests.post(url, headers=headers, data=data, timeout=60)
    if resp.status_code in (200, 201):
        return True, resp.status_code
    # try PUT (update) if POST rejected because object exists
    resp2 = requests.put(url, headers=headers, data=data, timeout=60)
    return (resp2.status_code in (200, 201)), f"{resp.status_code}/{resp2.status_code}: {resp2.text[:200]}"


def main():
    if not SERVICE_KEY:
        print("ERROR: SUPABASE_SERVICE_KEY not set", file=sys.stderr)
        sys.exit(2)
    ok = 0
    fail = []
    for doc_id, name, category in DOCS:
        data = build_doc(doc_id, name, category)
        success, info = upload(name, data)
        if success:
            ok += 1
            print(f"  OK  {name}  ({len(data)} bytes)")
        else:
            fail.append((name, info))
            print(f"  FAIL {name} -> {info}")
        time.sleep(0.05)
    print(f"\nDone: {ok}/{len(DOCS)} uploaded.")
    if fail:
        print("Failures:")
        for n, i in fail:
            print(f"  - {n}: {i}")
        sys.exit(1)


if __name__ == "__main__":
    main()
